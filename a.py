import streamlit as st
from PIL import Image
import easyocr
from docx import Document
from concurrent.futures import ThreadPoolExecutor

# Initialize EasyOCR Reader with GPU
reader = easyocr.Reader(['en'], gpu=True)

# Resize image to a maximum width
def resize_image(image, max_width=1024):
    width, height = image.size
    if width > max_width:
        new_height = int(max_width * height / width)
        return image.resize((max_width, new_height), Image.Resampling.LANCZOS)  # Use LANCZOS instead of ANTIALIAS
    return image

# Extract text function
def extract_text(image):
    return "\n".join(reader.readtext(image, detail=0))

# App Title
st.title("Fast Text Extractor from Images")
st.write("Upload images to extract text quickly using GPU acceleration.")

# Upload Section
uploaded_files = st.file_uploader("Upload Images", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

if uploaded_files:
    st.write(f"**{len(uploaded_files)} image(s) uploaded.**")
    document = Document()
    document.add_heading("Extracted Text", level=1)

    # Resize and Process Images in Parallel
    images = [resize_image(Image.open(uploaded_file)) for uploaded_file in uploaded_files]
    with ThreadPoolExecutor() as executor:
        texts = list(executor.map(extract_text, images))

    for uploaded_file, text in zip(uploaded_files, texts):
        st.image(uploaded_file, caption=uploaded_file.name, use_column_width=True)
        st.text_area(f"Extracted Text from {uploaded_file.name}", value=text, height=150)
        document.add_heading(f"Text from {uploaded_file.name}", level=2)
        document.add_paragraph(text)

    # Save the Word file
    output_file_path = "Extracted_Text.docx"
    document.save(output_file_path)

    # Provide a download link for the Word document
    with open(output_file_path, "rb") as file:
        st.download_button(
            label="Download Extracted Text as Word File",
            data=file,
            file_name="Extracted_Text.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
